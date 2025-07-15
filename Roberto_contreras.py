# Roberto Contreras - Selenium Script for Mercado Libre
import time
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from docx import Document
from docx.shared import Inches
import io

# --- Función para tomar Screenshot y agregarla a Word ---
def agregar_screenshot_a_word(step_title):
    """Toma una captura de pantalla y la agrega al documento de Word."""
    # Espera un momento para que la pantalla se renderice completamente
    time.sleep(1) 
    
    # Guarda la captura en un buffer de memoria para no crear archivos temporales
    image_stream = io.BytesIO(driver.get_screenshot_as_png())
    
    # Agrega un título para la captura en el documento
    document.add_heading(step_title, level=2)
    
    # Agrega la imagen desde el buffer de memoria
    document.add_picture(image_stream, width=Inches(6.0))
    document.add_page_break()
    print(f"📸 Captura del '{step_title}' guardada en el reporte.")

# Configurar navegador
options = webdriver.ChromeOptions()
options.add_argument('--start-maximized')
driver = webdriver.Chrome(options=options)

wait = WebDriverWait(driver, 10)

# --- Inicializar Documento de Word ---
document = Document()
document.add_heading('Reporte de Automatización - Mercado Libre', level=0)
print("📄 Creando reporte de Word...")

try:
    # 1. Entrar al sitio principal
    driver.get("https://www.mercadolibre.com/")
    
    agregar_screenshot_a_word('Paso 1: Página de Inicio de Mercado Libre México')   

    # 2. Seleccionar México
    mexico = wait.until(EC.element_to_be_clickable((By.XPATH, "//a[contains(@href, 'mercadolibre.com.mx')]")))
    mexico.click()

    agregar_screenshot_a_word('Paso 2:Seleccionar México')

    # 3. Buscar “playstation 5”
    search_input = wait.until(EC.presence_of_element_located((By.NAME, "as_word")))
    search_input.send_keys("playstation 5")
    search_input.send_keys(Keys.RETURN)

    overlay = wait.until(EC.element_to_be_clickable((By.XPATH, '/html/body/div[5]/div/div/div[2]/div/div/div[2]/button[2]/span')))
    overlay.click()

    agregar_screenshot_a_word('Paso 3: Buscar playstation 5')

    # 4. Filtro por condición “Nuevo”
    nuevo_filter = wait.until(EC.element_to_be_clickable((By.XPATH,'/html/body/main/div/div[2]/aside/section[2]/div[5]/ul/li[1]/a/span[1]')))
    driver.execute_script("arguments[0].scrollIntoView(true);", nuevo_filter)
    nuevo_filter.click()
    
    agregar_screenshot_a_word('Paso 4: Filtro por condición')
    
    # 5. filtrar por  ubicacion "Mexico City"
    localiza = wait.until(EC.element_to_be_clickable((By.XPATH,'//*[@id="root-app"]/div/div[2]/aside/section[2]/div[14]/ul/li[1]/a/span[1]')))
    driver.execute_script("arguments[0].scrollIntoView(true);", localiza)
    localiza.click()
    #time.sleep(1000)
    agregar_screenshot_a_word('Paso 5: Filtro por ubicación')

    # 6. Ordenar por “Mayor precio”
    sort_button = wait.until(EC.element_to_be_clickable((By.XPATH,'//*[@id=":R1b55ie:-display-values"]')))
    #driver.execute_script("arguments[0].scrollIntoView(true);", nuevo_filter)
    sort_button.click()

    mayor_precio_option = wait.until(EC.element_to_be_clickable((By.XPATH,'//*[@id=":R1b55ie:-menu-list-option-price_desc"]/div/div/span')))
    mayor_precio_option.click()
    #time.sleep(1000)

    # 7. Get name and price of the first 5 products
    # time.sleep(3)  # Wait for results to load
    # products = driver.find_elements(By.CSS_SELECTOR, ".ui-search-result__wrapper")[:5]

    # print("\n🔍 Primeros 5 productos encontrados:\n")
    # for i, product in enumerate(products, 1):
    #     try:
    #         name = product.find_element(By.CSS_SELECTOR, "h2").text
    #         price_elements = product.find_elements(By.CSS_SELECTOR, ".ui-search-price__part")
    #         if price_elements:
    #             price = " ".join([elem.text for elem in price_elements if elem.text.strip()])
    #         else:
    #             price = "Price not available"
    #         print(f"{i}. {name} - {price}")
    #     except:
    #         print(f"{i}.Product missing complete data")
    document.add_heading('Paso 7: Extracción de Datos de Productos', level=2)
    print("\n🔍 Primeros 5 productos encontrados:\n")
    
    # Se agrega el texto de los productos al documento de Word
    parrafo_productos = document.add_paragraph()
    parrafo_productos.add_run('Primeros 5 productos encontrados:\n').bold = True

    wait.until(EC.presence_of_element_located((By.CSS_SELECTOR, "li.ui-search-layout__item")))
    time.sleep(2) 

    products = driver.find_elements(By.CSS_SELECTOR, "li.ui-search-layout__item")[:5]

    for i, product in enumerate(products, 1):
        try:
            name = product.find_element(By.CSS_SELECTOR, "h2.ui-search-item__title").text
            price = product.find_element(By.CSS_SELECTOR, "span.andes-money-amount__fraction").text
            product_info = f"{i}. {name} - ${price}\n"
            print(product_info, end="")
            parrafo_productos.add_run(product_info) # Agrega la info al párrafo
        except Exception:
            product_info = f"{i}. No se pudieron obtener los datos completos del producto.\n"
            print(product_info, end="")
            parrafo_productos.add_run(product_info)
except Exception as e:
    print(f"❌ Ocurrió un error: {e}")
finally:
    time.sleep(5)
    document.save('reporte.docx')
    driver.quit()

